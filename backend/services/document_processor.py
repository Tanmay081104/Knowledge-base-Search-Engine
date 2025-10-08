"""Document processing service for extracting and chunking text from various file formats."""

import os
import uuid
import json
import re
from typing import List, Dict, Any
from pathlib import Path
import PyPDF2
import aiofiles
from loguru import logger
import docx
import pandas as pd
import openpyxl
from PIL import Image
import pytesseract
import io

from backend.app.config import settings

class DocumentProcessor:
    """Service for processing and chunking documents."""
    
    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.allowed_extensions = settings.allowed_extensions
    
    async def process_uploaded_file(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """
        Process an uploaded file and extract text content.
        
        Args:
            file_path: Path to the uploaded file
            original_filename: Original filename
            
        Returns:
            Dictionary containing document info and extracted text
        """
        try:
            # Generate unique document ID
            document_id = str(uuid.uuid4())
            
            # Extract file extension
            file_extension = Path(original_filename).suffix.lower().lstrip('.')
            
            if file_extension not in self.allowed_extensions:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Extract text based on file type
            if file_extension == 'pdf':
                text_content = await self._extract_pdf_text(file_path)
            elif file_extension in ['txt', 'md']:
                text_content = await self._extract_text_file(file_path)
            elif file_extension == 'docx':
                text_content = await self._extract_docx_text(file_path)
            elif file_extension in ['xlsx', 'xls']:
                text_content = await self._extract_excel_text(file_path)
            elif file_extension in ['jpg', 'jpeg', 'png', 'bmp', 'tiff']:
                text_content = await self._extract_image_text(file_path)
            elif file_extension == 'json':
                text_content = await self._extract_json_text(file_path)
            else:
                raise ValueError(f"Processing for {file_extension} not implemented yet")
            
            # Create text chunks
            chunks = self._create_text_chunks(text_content)
            
            # Get file size
            file_size = os.path.getsize(file_path)
            
            result = {
                'document_id': document_id,
                'filename': original_filename,
                'file_size': file_size,
                'text_content': text_content,
                'chunks': chunks,
                'chunks_count': len(chunks)
            }
            
            logger.info(f"Processed document {original_filename} into {len(chunks)} chunks")
            return result
            
        except Exception as e:
            logger.error(f"Error processing file {original_filename}: {str(e)}")
            raise
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text_content = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            async with aiofiles.open(file_path, mode='r', encoding='utf-8') as file:
                content = await file.read()
                return content.strip()
                
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, mode='r', encoding='latin-1') as file:
                    content = await file.read()
                    return content.strip()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {str(e)}")
                raise ValueError(f"Failed to read text file: {str(e)}")
        
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text: {str(e)}")
    
    def _create_text_chunks(self, text: str) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: The text to chunk
            
        Returns:
            List of chunk dictionaries
        """
        if not text.strip():
            return []
        
        # Simple sentence-aware chunking
        sentences = self._split_into_sentences(text)
        chunks = []
        current_chunk = ""
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            # If adding this sentence would exceed chunk size and we have content
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # Create chunk
                chunk_id = str(uuid.uuid4())
                chunks.append({
                    'chunk_id': chunk_id,
                    'content': current_chunk.strip(),
                    'length': len(current_chunk.strip())
                })
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk, self.chunk_overlap)
                current_chunk = overlap_text + " " + sentence if overlap_text else sentence
                current_length = len(current_chunk)
            else:
                # Add sentence to current chunk
                if current_chunk:
                    current_chunk += " " + sentence
                else:
                    current_chunk = sentence
                current_length = len(current_chunk)
        
        # Add final chunk if it has content
        if current_chunk.strip():
            chunk_id = str(uuid.uuid4())
            chunks.append({
                'chunk_id': chunk_id,
                'content': current_chunk.strip(),
                'length': len(current_chunk.strip())
            })
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        import re
        
        # Simple sentence splitting - can be improved with NLTK
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        return sentences
    
    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """Get the last part of text for overlap."""
        if len(text) <= overlap_size:
            return text
        
        # Try to find a sentence boundary within overlap size
        overlap_text = text[-overlap_size:]
        sentence_start = overlap_text.find('. ')
        
        if sentence_start != -1:
            return overlap_text[sentence_start + 2:]
        
        return overlap_text
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file with advanced formatting preservation."""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract paragraphs with formatting info
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                text_content.append("\n".join(table_text))
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_excel_text(self, file_path: str) -> str:
        """Extract text from Excel file with smart data interpretation."""
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            text_content = []
            
            for sheet_name, df in excel_data.items():
                text_content.append(f"Sheet: {sheet_name}")
                
                # Convert DataFrame to text representation
                if not df.empty:
                    # Add column headers
                    headers = " | ".join(str(col) for col in df.columns)
                    text_content.append(f"Headers: {headers}")
                    
                    # Add data rows (limit to prevent huge texts)
                    for idx, row in df.head(100).iterrows():
                        row_text = " | ".join(str(val) for val in row.values)
                        text_content.append(row_text)
                
                text_content.append("\n")
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error extracting text from Excel {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from Excel: {str(e)}")
    
    async def _extract_image_text(self, file_path: str) -> str:
        """Extract text from image using OCR (Optical Character Recognition)."""
        try:
            # Open image
            image = Image.open(file_path)
            
            # Apply OCR
            text_content = pytesseract.image_to_string(image, config='--oem 3 --psm 6')
            
            if not text_content.strip():
                return "No text detected in the image."
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from image {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from image: {str(e)}")
    
    async def _extract_json_text(self, file_path: str) -> str:
        """Extract and format text from JSON file."""
        try:
            async with aiofiles.open(file_path, mode='r', encoding='utf-8') as file:
                content = await file.read()
                json_data = json.loads(content)
            
            # Convert JSON to readable text
            text_content = self._json_to_text(json_data)
            return text_content
            
        except Exception as e:
            logger.error(f"Error extracting text from JSON {file_path}: {str(e)}")
            raise ValueError(f"Failed to extract text from JSON: {str(e)}")
    
    def _json_to_text(self, data, prefix="") -> str:
        """Recursively convert JSON data to readable text."""
        text_parts = []
        
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    text_parts.append(f"{prefix}{key}:")
                    text_parts.append(self._json_to_text(value, prefix + "  "))
                else:
                    text_parts.append(f"{prefix}{key}: {value}")
        elif isinstance(data, list):
            for i, item in enumerate(data):
                text_parts.append(f"{prefix}[{i}]:")
                text_parts.append(self._json_to_text(item, prefix + "  "))
        else:
            text_parts.append(f"{prefix}{data}")
        
        return "\n".join(text_parts)
    
    def validate_file(self, filename: str, file_size: int) -> bool:
        """
        Validate uploaded file.
        
        Args:
            filename: Name of the file
            file_size: Size of the file in bytes
            
        Returns:
            True if valid, False otherwise
        """
        # Check file extension
        file_extension = Path(filename).suffix.lower().lstrip('.')
        if file_extension not in self.allowed_extensions:
            logger.warning(f"Invalid file extension: {file_extension}")
            return False
        
        # Check file size
        if file_size > settings.max_file_size:
            logger.warning(f"File too large: {file_size} bytes")
            return False
        
        return True
